import os
import logging
import threading
import queue
import time
import uuid
import shutil
import sys
from PIL import Image
from docx import Document
from pdf2docx import Converter as PDFConverter

# Conditional import for Windows-only COM
HAS_COM = False
if os.name == 'nt':
    try:
        import comtypes.client
        HAS_COM = True
    except ImportError:
        pass

# Fallback libraries for Linux
HAS_ASPOSE = False
try:
    import aspose.words as aw
    import aspose.cells as ac
    HAS_ASPOSE = True
except ImportError:
    pass

logger = logging.getLogger(__name__)

# Queue for conversion tasks
_conversion_queue = queue.Queue()

# Dictionary to store results
_results = {}

def _safe_remove(path):
    if os.path.exists(path):
        for _ in range(5):
            try:
                os.remove(path)
                return
            except:
                time.sleep(0.2)

def _conversion_worker():
    logger.info("Starting background conversion worker...")
    if HAS_COM:
        comtypes.CoInitialize()
    
    word = None
    excel = None
    
    if HAS_COM:
        try:
            word = comtypes.client.CreateObject('Word.Application')
            word.Visible = False
            word.DisplayAlerts = 0 # wdAlertsNone
            excel = comtypes.client.CreateObject('Excel.Application')
            excel.Visible = False
            excel.DisplayAlerts = False
            logger.info("Office applications pre-launched successfully.")
        except Exception as e:
            logger.warning(f"Could not pre-launch Office apps: {e}")

    while True:
        task = _conversion_queue.get()
        if task is None: break
            
        task_id, input_path, output_path, ext, target_format = task
        
        try:
            abs_out = os.path.abspath(output_path)
            
            # Special Case: Image Merging
            if ext == "images_merge":
                input_paths = input_path # It's a list in this case
                if target_format == 'pdf':
                    processed_imgs = []
                    for p in input_paths:
                        img = Image.open(p)
                        if img.mode != 'RGB': img = img.convert('RGB')
                        processed_imgs.append(img)
                    
                    if processed_imgs:
                        processed_imgs[0].save(abs_out, save_all=True, append_images=processed_imgs[1:])
                        for img in processed_imgs: img.close()
                
                elif target_format == 'docx':
                    doc_obj = Document()
                    for p in input_paths:
                        doc_obj.add_picture(os.path.abspath(p))
                    doc_obj.save(abs_out)
                
                _results[task_id] = {"status": "success"}
                continue # Use continue instead of return to process next task

            abs_in = os.path.abspath(input_path)

            if target_format == 'pdf':
                if ext in ['.docx', '.doc']:
                    if HAS_COM:
                        if word is None: word = comtypes.client.CreateObject('Word.Application'); word.Visible = False; word.DisplayAlerts = 0
                        doc = word.Documents.Open(abs_in, ReadOnly=True)
                        try:
                            doc.SaveAs(abs_out, FileFormat=17) # wdFormatPDF
                        finally:
                            doc.Close(0) # wdDoNotSaveChanges
                    elif HAS_ASPOSE:
                        doc = aw.Document(abs_in)
                        doc.save(abs_out)
                    else:
                        raise ImportError("No conversion engine available for Word to PDF (Windows COM or Aspose required)")

                elif ext in ['.xlsx', '.xls']:
                    if HAS_COM:
                        if excel is None: excel = comtypes.client.CreateObject('Excel.Application'); excel.Visible = False; excel.DisplayAlerts = False
                        wb = excel.Workbooks.Open(abs_in, ReadOnly=True)
                        try:
                            wb.ExportAsFixedFormat(0, abs_out) # xlTypePDF
                        finally:
                            wb.Close(False)
                    elif HAS_ASPOSE:
                        workbook = ac.Workbook(abs_in)
                        workbook.save(abs_out, ac.SaveFormat.PDF)
                    else:
                        raise ImportError("No conversion engine available for Excel to PDF (Windows COM or Aspose required)")

                elif ext in ['.png', '.jpg', '.jpeg']:
                    img = Image.open(abs_in)
                    try:
                        if img.mode != 'RGB': img = img.convert('RGB')
                        img.save(abs_out, "PDF", resolution=100.0)
                    finally:
                        img.close()
                elif ext == '.txt':
                    temp_docx = input_path + ".docx"
                    doc_obj = Document()
                    with open(abs_in, 'r', encoding='utf-8', errors='ignore') as f:
                        for line in f: doc_obj.add_paragraph(line.strip())
                    doc_obj.save(os.path.abspath(temp_docx))
                    
                    if HAS_COM:
                        if word is None: word = comtypes.client.CreateObject('Word.Application'); word.Visible = False; word.DisplayAlerts = 0
                        wdoc = word.Documents.Open(os.path.abspath(temp_docx), ReadOnly=True)
                        try:
                            wdoc.SaveAs(abs_out, FileFormat=17)
                        finally:
                            wdoc.Close(0)
                    elif HAS_ASPOSE:
                        doc = aw.Document(os.path.abspath(temp_docx))
                        doc.save(abs_out)
                    
                    _safe_remove(temp_docx)
                elif ext == '.pdf':
                    shutil.copy2(abs_in, abs_out)
                else:
                    raise ValueError(f"No PDF conversion path for {ext}")

            elif target_format == 'docx':
                if ext in ['.docx', '.doc']:
                    if ext == '.docx':
                        shutil.copy2(abs_in, abs_out)
                    else: # .doc to .docx
                        if HAS_COM:
                            if word is None: word = comtypes.client.CreateObject('Word.Application'); word.Visible = False; word.DisplayAlerts = 0
                            doc = word.Documents.Open(abs_in, ReadOnly=True)
                            try:
                                doc.SaveAs(abs_out, FileFormat=16) # wdFormatXMLDocument
                            finally:
                                doc.Close(0)
                        elif HAS_ASPOSE:
                            doc = aw.Document(abs_in)
                            doc.save(abs_out)
                elif ext == '.pdf':
                    cv = PDFConverter(abs_in)
                    try:
                        cv.convert(abs_out)
                    finally:
                        cv.close()
                elif ext == '.txt':
                    doc_obj = Document()
                    with open(abs_in, 'r', encoding='utf-8', errors='ignore') as f:
                        for line in f: doc_obj.add_paragraph(line.strip())
                    doc_obj.save(abs_out)
                elif ext in ['.png', '.jpg', '.jpeg']:
                    doc_obj = Document()
                    doc_obj.add_picture(abs_in)
                    doc_obj.save(abs_out)
                else:
                    raise ValueError(f"No Word conversion path for {ext}")

            _results[task_id] = {"status": "success"}
        except Exception as e:
            logger.error(f"Conversion error for {task_id}: {e}")
            _results[task_id] = {"status": "error", "message": str(e)}
        finally:
            _conversion_queue.task_done()

    if word:
        try: word.Quit()
        except: pass
    if excel:
        try: excel.Quit()
        except: pass
    if HAS_COM:
        comtypes.CoUninitialize()

_worker_thread = threading.Thread(target=_conversion_worker, daemon=True)
_worker_thread.start()

def convert_to_pdf(input_path: str, output_path: str, ext: str, target_format: str = "pdf"):
    task_id = str(uuid.uuid4())
    _results[task_id] = {"status": "pending"}
    _conversion_queue.put((task_id, input_path, output_path, ext, target_format))
    while True:
        res = _results.get(task_id)
        if res["status"] == "success":
            del _results[task_id]
            return
        elif res["status"] == "error":
            err_msg = res["message"]
            del _results[task_id]
            raise Exception(err_msg)
        time.sleep(0.1)
